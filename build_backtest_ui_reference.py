"""
build_backtest_ui_reference.py
Generates Backtest_UI_Reference.docx — FVB33 Backtest Tab UI Reference Guide.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\mohin\Documents\PythonCodes\get_trading_data\Backtest_UI_Reference.docx"

# ── Colours ────────────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
ALT_BLUE    = RGBColor(0xD6, 0xE4, 0xF0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
MONO_COLOR  = RGBColor(0x1A, 0x1A, 0x2E)
NOTE_BG     = "FFF9C4"   # light yellow
CODE_BG     = "F2F2F2"   # light gray

FONT_BODY = "Calibri"
FONT_MONO = "Courier New"


# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    for e in tcPr.findall(qn("w:shd")):
        tcPr.remove(e)
    tcPr.append(shd)


def _set_cell_borders(cell, color="BBBBBB", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    for e in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(e)
    tcPr.append(borders)


def _set_col_widths(table, widths_cm):
    tbl = table._tbl
    for g in tbl.findall(qn("w:tblGrid")):
        tbl.remove(g)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w / 2.54 * 1440)))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _set_cell_width(cell, w_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(w_cm / 2.54 * 1440)))
    tcW.set(qn("w:type"), "dxa")
    for e in tcPr.findall(qn("w:tcW")):
        tcPr.remove(e)
    tcPr.append(tcW)


def _cell_para(cell, text, bold=False, mono=False, size=9,
               color=None, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = FONT_MONO if mono else FONT_BODY
    r.font.color.rgb = color or (WHITE if bold and not mono else (MONO_COLOR if mono else BLACK))


# ── Table builder ──────────────────────────────────────────────────────────────

def make_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    _set_col_widths(table, widths_cm)

    # Header
    for ci, (cell, hdr) in enumerate(zip(table.rows[0].cells, headers)):
        _set_cell_bg(cell, "1F497D")
        _set_cell_borders(cell, "1F497D")
        _set_cell_width(cell, widths_cm[ci])
        _cell_para(cell, hdr, bold=True, size=9, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        fill = "D6E4F0" if ri % 2 == 0 else "FFFFFF"
        for ci, (cell, val) in enumerate(zip(table.rows[ri + 1].cells, row_data)):
            _set_cell_bg(cell, fill)
            _set_cell_borders(cell, "BBBBBB", 4)
            _set_cell_width(cell, widths_cm[ci])
            _cell_para(cell, val, bold=(ci == 0), mono=(ci == 0), size=9)

    doc.add_paragraph()


# ── Text helpers ───────────────────────────────────────────────────────────────

def set_margins(doc, top=1.8, bottom=1.8, left=2.0, right=2.0):
    for s in doc.sections:
        s.top_margin    = Cm(top)
        s.bottom_margin = Cm(bottom)
        s.left_margin   = Cm(left)
        s.right_margin  = Cm(right)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = FONT_BODY
    r.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = FONT_BODY
    r.font.color.rgb = DARK_BLUE


def add_body(doc, text, size=10, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT_BODY
    r.italic = italic
    r.font.color.rgb = color or RGBColor(0x22, 0x22, 0x22)


def add_note(doc, text, bg_hex=NOTE_BG):
    """Shaded note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg_hex)
    pPr.append(shd)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.name = FONT_BODY
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x00)


def add_code_block(doc, lines):
    """Light-gray monospace block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  CODE_BG)
        pPr.append(shd)
        r = p.add_run(line)
        r.font.size = Pt(8.5)
        r.font.name = FONT_MONO
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()


def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT_BODY
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_table_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = FONT_BODY
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Page ")
        r.font.size = Pt(9)
        r.font.name = FONT_BODY
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        from docx.oxml import OxmlElement
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = FONT_BODY
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc)
    add_footer(doc)

    # ── Title ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("FVB33 Backtest Tab")
    r.bold = True; r.font.size = Pt(22); r.font.name = FONT_BODY
    r.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("UI Reference Guide")
    r2.bold = True; r2.font.size = Pt(16); r2.font.name = FONT_BODY
    r2.font.color.rgb = DARK_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run("Signal Logic, Column Definitions & Calculations")
    r3.italic = True; r3.font.size = Pt(11); r3.font.name = FONT_BODY
    r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(24)
    r4 = p4.add_run("April 2026")
    r4.italic = True; r4.font.size = Pt(10); r4.font.name = FONT_BODY
    r4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ── Section 1 ─────────────────────────────────────────────────────────────
    add_h1(doc, "1.  Signal Columns — entry_signal vs sig_entry")
    add_body(doc,
        "There are two different signal columns in the data pipeline. "
        "Understanding which one the Backtest tab uses is critical to interpreting results correctly.")

    add_table_label(doc, "Signal Column Comparison")
    make_table(doc,
        headers=["Column", "Table", "BX Condition", "Price Buffer", "State Machine"],
        widths_cm=[3.2, 2.2, 7.0, 3.3, 3.8],
        rows=[
            ["entry_signal", "Daily + Weekly",
             "bx_monthly_dark_red==0 (same as sig_entry v1.3)",
             "None — Close must be AT or below upper band",
             "No — simple row-by-row calculation"],
            ["sig_entry (v1.3)", "Weekly only",
             "bx_monthly_dark_red==0  (0=green or recovering; 1=dark red=no entry)",
             "+3% upper / -25% lower buffer",
             "Yes — only fires from FLAT_SL state"],
        ]
    )

    add_note(doc,
        "NOTE: The Backtest tab uses sig_entry (not entry_signal). "
        "Both now use bx_monthly_dark_red==0 as the BX gate (v1.3). "
        "dark_red==0 means the last COMPLETED monthly candle is either positive (green) "
        "OR negative but recovering vs the prior month. "
        "Entry is blocked ONLY when dark_red==1 (last closed month is negative AND still falling). "
        "Key remaining difference: sig_entry has a +3%/-25% price buffer and uses a state machine.")

    add_table_label(doc,
        "Signal version history — entry condition evolution:")
    make_table(doc,
        headers=["Version", "Entry BX Condition", "Status / Effect"],
        widths_cm=[2.0, 7.5, 10.0],
        rows=[
            ["v1.1",
             "bx_monthly_dark_red==0",
             "BUG: indicators.py treated the last row of the dataset as a month-end, "
             "causing partial current-month BX to leak into bx_monthly_completed. "
             "Dark red could flip to 0 mid-month incorrectly."],
            ["v1.2",
             "bx_monthly_completed > 0",
             "Too strict: blocked entry when monthly BX was negative but recovering. "
             "e.g. BJRI at -1.78 (recovering from -5.25) was incorrectly blocked."],
            ["v1.3 (current)",
             "bx_monthly_dark_red==0  (is_month_end bug fixed)",
             "CORRECT. Entry allowed when last closed month is green OR recovering. "
             "Blocked only when dark red (negative AND still falling). "
             "indicators.py fixed: last dataset row no longer treated as a month-end."],
        ]
    )

    # ── Section 2 ─────────────────────────────────────────────────────────────
    add_h1(doc, "2.  Snapshot Data Loading — load_backtest_snapshot()")
    add_body(doc,
        "The Backtest tab loads the weekly snapshot (one row per ticker, latest week only) "
        "from Delta Lake, then adds derived columns at query time. "
        "These derived columns are NOT stored in the Delta table.")

    add_table_label(doc, "Derived Columns Added at Query Time")
    make_table(doc,
        headers=["Column", "Formula", "Notes"],
        widths_cm=[3.8, 7.0, 6.7],
        rows=[
            ["_dist_upper_pct",
             "(Close - fvb_upper_thresh) / fvb_upper_thresh x 100",
             "Negative = price is below upper band (in zone). Positive = above band."],
            ["_upside_to_tp_pct",
             "(tpx1 - Close) / Close x 100",
             "How much % upside remains to the TPX1 target from current price."],
            ["_unrealised_pct",
             "(Close - sig_entry_price) / sig_entry_price x 100",
             "Only populated when sig_in_position==1. NaN otherwise."],
            ["_fvb33_label",
             '"Bullish" if fvb_band_green==1, else "Bearish"',
             "Human-readable label for FVB33 band state."],
            ["_bx_monthly_label",
             '"Bullish" if bx_monthly_completed > 0; "Bearish" if <= 0',
             "Based on last CLOSED monthly candle. Null shown as dash."],
            ["_bx_weekly_label",
             '"Bullish" if bx > 0, else "Bearish"',
             "Based on chart-period (weekly) BX value."],
            ["_bx_mo_wk",
             "Combination of monthly + weekly labels",
             'Values: "Bull / Bull", "Bull / Bear", "Bear / Bull", "Bear / Bear"'],
            ["sector",
             "Joined from SQLite universe DB via symbol lookup",
             "Dash for SPACs and warrants. Sourced from NASDAQ screener bulk download."],
        ]
    )

    # ── Section 3 ─────────────────────────────────────────────────────────────
    add_h1(doc, "3.  Tab 1 — Active Signals")
    add_body(doc,
        "Shows all tickers with a signal this week. "
        "The user selects which signal types to display via a multiselect filter.")

    add_h2(doc, "Signal Filter Logic")
    make_table(doc,
        headers=["Label", "Column Checked", "Meaning"],
        widths_cm=[2.5, 4.0, 11.0],
        rows=[
            ["Entry",    "sig_entry == 1",    "New position opened from FLAT_SL state this week."],
            ["Re-Entry", "sig_reentry == 1",  "Re-opened after a prior TP; same entry conditions, from FLAT_TP state."],
            ["Add",      "sig_add == 1",      "Price pulled back to FVB33 entry zone while already LONG — add to position."],
            ["TP Hit",   "sig_tp == 1",       "Weekly High >= TPX1 this week — trade closed in profit at target."],
            ["SL",       "sig_sl == 1",       "bx_monthly_dark_red==1 this week — trade stopped out."],
        ]
    )

    add_h2(doc, "Additional Filters")
    make_table(doc,
        headers=["Filter", "Column", "Logic"],
        widths_cm=[3.0, 4.0, 10.5],
        rows=[
            ["FVB33",      "_fvb33_label",      'Equality match: "Bullish" or "Bearish"'],
            ["Monthly BX", "_bx_monthly_label", 'Equality match: "Bullish" or "Bearish"'],
            ["Mo/Wk BX",   "_bx_mo_wk",         'Equality match: "Bull/Bull", "Bull/Bear", "Bear/Bull", "Bear/Bear"'],
            ["Sector",     "sector",             "Multi-select — any of the selected sectors."],
            ["Min Price",  "close",              "close >= min_price threshold (default $5)."],
        ]
    )

    add_h2(doc, "Displayed Columns")
    make_table(doc,
        headers=["Column", "Source", "Calculation / Notes"],
        widths_cm=[3.5, 4.5, 9.5],
        rows=[
            ["Symbol",        "symbol",          "Direct from snapshot."],
            ["Sector",        "sector",          "Joined from universe DB."],
            ["Signal",        "sig_entry / sig_reentry / sig_add / sig_tp / sig_sl",
             "Badge built from whichever signal flags are 1 on this row."],
            ["Price",         "close",           "Current week closing price."],
            ["Dist FVB33%",   "_dist_upper_pct",
             "(Close - fvb_upper_thresh) / fvb_upper_thresh x 100. Negative = below band."],
            ["Upside to TP",  "_upside_to_tp_pct", "(tpx1 - Close) / Close x 100"],
            ["FVB33",         "_fvb33_label",    "Bullish / Bearish based on fvb_band_green."],
            ["Monthly BX",    "_bx_monthly_label", "Bullish / Bearish based on bx_monthly_completed sign."],
            ["Mo/Wk BX",      "_bx_mo_wk",       "Combined monthly + weekly BX label."],
            ["Cycle",         "cycle_bull, cycle_months, cycle_pct",
             "Icon + months elapsed + % gain since current cycle started."],
            ["Win Rate",      "bt_win_rate",     "% of historical trades with pct_gain > 0."],
            ["Profit Factor", "bt_profit_factor","sum(winning pcts) / abs(sum(losing pcts))."],
            ["Geo Return",    "bt_geo_return",   "Compounded return across all historical trades."],
            ["# Trades",      "bt_trades",       "Count of completed trades in full history."],
        ]
    )

    # ── Section 4 ─────────────────────────────────────────────────────────────
    add_h1(doc, "4.  Tab 2 — In Position")
    add_body(doc,
        "Filters snapshot to sig_in_position == 1. "
        "Shows all currently held positions with live unrealised P&L.")

    add_note(doc,
        "Summary bar: total positions, count profitable (unrealised > 0), "
        "average unrealised %, and best performing position.")

    add_table_label(doc, "Displayed Columns")
    make_table(doc,
        headers=["Column", "Source", "Calculation / Notes"],
        widths_cm=[3.5, 4.5, 9.5],
        rows=[
            ["Symbol",       "symbol",          "Direct."],
            ["Sector",       "sector",          "Universe DB join."],
            ["Price",        "close",           "Current week closing price."],
            ["Entry Price",  "sig_entry_price",
             "Close on the week the trade was opened; carried forward each week while in position."],
            ["Unrealised",   "_unrealised_pct",
             "(Close - sig_entry_price) / sig_entry_price x 100"],
            ["TPX1",         "tpx1",
             "The weekly TP target price — SMA(OHLC4,20) + 4.3 x EMA(TrueRange,66)."],
            ["Upside to TP", "_upside_to_tp_pct", "(tpx1 - Close) / Close x 100"],
            ["FVB33",        "_fvb33_label",    "Bullish / Bearish."],
            ["Monthly BX",   "_bx_monthly_label", "Bullish / Bearish."],
            ["Mo/Wk BX",     "_bx_mo_wk",       "Bull/Bull, Bull/Bear, etc."],
            ["Cycle",        "cycle_bull, cycle_months, cycle_pct",
             "Bull/bear status + duration + gain since cycle started."],
            ["Win Rate",     "bt_win_rate",     "Historical win rate for this ticker."],
            ["Profit Factor","bt_profit_factor","Historical profit factor."],
            ["Wks In Trade", "bt_weeks_in_trade",
             "Total weeks sig_in_position==1 across ALL history — not just the current trade."],
        ]
    )

    # ── Section 5 ─────────────────────────────────────────────────────────────
    add_h1(doc, "5.  Tab 3 — Leaderboard")
    add_body(doc,
        "Ranks all symbols by backtest performance. "
        "All stats come from bt_* columns pre-computed by signals.py "
        "and broadcast to every row of the symbol in the Delta weekly table.")

    add_h2(doc, "Filter Controls")
    make_table(doc,
        headers=["Control", "Column", "Default"],
        widths_cm=[4.5, 4.5, 8.5],
        rows=[
            ["Min Trades",       "bt_trades",        "5"],
            ["Min Win Rate %",   "bt_win_rate",       "0%"],
            ["Min Profit Factor","bt_profit_factor",  "1.0"],
            ["Sector",           "sector",            "All"],
        ]
    )

    add_h2(doc, "All Backtest Stat Columns (bt_*)")
    make_table(doc,
        headers=["Column", "Formula", "Notes"],
        widths_cm=[4.2, 6.5, 6.8],
        rows=[
            ["bt_trades",          "Count of completed TP + SL exits",
             "Open (unfinished) trades are excluded from all stats."],
            ["bt_tp_count",        "Count of trades where sig_tp fired",
             "Trade exited at TPX1 — High >= tpx1."],
            ["bt_sl_count",        "Count of trades where sig_sl fired",
             "Trade exited at Close when bx_monthly_dark_red==1."],
            ["bt_reentry_count",   "Count of trades entered via sig_reentry",
             "Re-entry after a TP exit."],
            ["bt_win_rate",        "sum(pct_gain > 0) / bt_trades x 100",
             "WIN = trade closed in profit, regardless of TP or SL exit type."],
            ["bt_avg_tp_pct",      "mean(pct_gain) for TP-exit trades",
             "TP exit price = tpx1 value on the TP bar."],
            ["bt_avg_sl_pct",      "mean(pct_gain) for SL-exit trades",
             "SL exit price = Close on the SL bar. Usually negative."],
            ["bt_profit_factor",   "sum(pct > 0) / abs(sum(pct <= 0))",
             "Profit-based — not exit-type-based. Higher = better."],
            ["bt_geo_return",      "(product(1 + pct/100) - 1) x 100",
             "e.g. +20% then -10% then +15% = 1.2 x 0.9 x 1.15 = 1.242 = +24.2%"],
            ["bt_reentry_win_rate","sum(pct > 0 for re-entries) / bt_reentry_count x 100",
             "Win rate specifically for re-entry trades."],
            ["bt_weeks_in_trade",  "sum(sig_in_position) across full weekly history",
             "Total weeks ever held for this ticker."],
        ]
    )

    # ── Section 6 ─────────────────────────────────────────────────────────────
    add_h1(doc, "6.  Tab 4 — Deep Dive")
    add_body(doc,
        "Single-symbol analysis. Calls load_symbol_trades(ticker) which re-reads the FULL weekly "
        "history (all bars, not just the latest snapshot row) and reconstructs every trade live "
        "from the signal columns.")

    add_h2(doc, "Trade Reconstruction Logic")
    add_code_block(doc, [
        "For each weekly bar in chronological order:",
        "",
        "  IF sig_entry==1 OR sig_reentry==1:",
        "      entry_price = Close of that bar",
        "      entry_date  = week_start of that bar",
        "",
        "  IF sig_tp==1 AND entry_price is set:",
        "      exit_price = tpx1 value on that bar",
        "      pct_gain   = (tpx1 - entry_price) / entry_price x 100",
        "      hold_weeks = (exit_date - entry_date).days / 7",
        "      --> record trade, clear entry_price",
        "",
        "  IF sig_sl==1 AND entry_price is set:",
        "      exit_price = Close of that bar",
        "      pct_gain   = (Close - entry_price) / entry_price x 100",
        "      hold_weeks = (exit_date - entry_date).days / 7",
        "      --> record trade, clear entry_price",
    ])

    add_h2(doc, "Row 1 — Strategy Performance Metrics")
    make_table(doc,
        headers=["Metric", "Formula", "Notes"],
        widths_cm=[3.5, 6.5, 7.5],
        rows=[
            ["Trades",        "len(trades)",
             "Count of completed trades (TP + SL exits)."],
            ["Accuracy",      "sum(pct_gain > 0) / len(trades) x 100",
             "% of trades that closed in profit. WIN = pct_gain > 0 regardless of TP or SL."],
            ["Avg TP Gain",   "mean(pct_gain) for TP-exit trades only",
             "Average % gain when the trade hit TPX1."],
            ["Avg SL Loss",   "mean(pct_gain) for SL-exit trades only",
             "Average % gain/loss when stopped out. Usually negative."],
            ["Profit Factor", "sum(pct > 0) / abs(sum(pct <= 0))",
             "Across ALL trades. > 1.0 means strategy is profitable overall."],
            ["Geo Return",    "(product(1 + pct/100) - 1) x 100",
             "Compounded return if you followed every signal. Does NOT assume position sizing."],
        ]
    )

    add_h2(doc, "Row 2 — Risk & Cycle Metrics")
    make_table(doc,
        headers=["Metric", "Formula", "Notes"],
        widths_cm=[3.5, 6.5, 7.5],
        rows=[
            ["Risk : Reward",
             "abs(avg_tp_gain / avg_sl_loss) shown as '1 : X.X'",
             "e.g. avg TP = +20%, avg SL = -8%  ->  Risk:Reward = 1 : 2.5"],
            ["Weeks In Trade",
             "sum(sig_in_position) across full weekly history",
             "Total weeks ever held. Includes current open trade if any."],
            ["Avg Hold",
             "mean(hold_weeks) across all completed trades",
             "Average duration per trade in weeks."],
            ["Avg Bull Pullback",
             "cycle_avg_pullback from snapshot",
             "Average max drawdown within bull cycles before price recovers. "
             "Sourced from cycle stats, not from trade list."],
        ]
    )

    add_h2(doc, "Current Status Section")
    add_body(doc,
        "Pulled from the snapshot row for the selected ticker. Shows: current price, FVB33 state, "
        "Monthly BX, Mo/Wk BX combined, and whether currently In Position. "
        "If in position, also shows: Entry Price, Unrealised P&L%, and TPX1 target.")

    add_h2(doc, "Equity Curve")
    add_body(doc,
        "Starts at $100 and compounds each trade sequentially: "
        "equity[i+1] = equity[i] x (1 + pct_gain / 100). "
        "Visualised as a line chart indexed by trade number.")

    add_h2(doc, "Price Chart (last 5 years)")
    add_body(doc,
        "Line chart of weekly Close, FVB33 Upper Band (fvb_upper_thresh), "
        "FVB33 Lower Band (fvb_lower_thresh), and TPX1 target — all from the full weekly history, "
        "filtered to the last 5 years.")

    add_h2(doc, "Trade History Table")
    add_body(doc,
        "One row per completed trade showing: entry date, exit date, "
        "entry type (Entry or Re-Entry), exit type (TP or SL), "
        "entry price, exit price, P&L%, and weeks held.")

    # ── Section 7 ─────────────────────────────────────────────────────────────
    add_h1(doc, "7.  Important Notes")

    add_bullet(doc,
        "Cache: The snapshot is cached for 300 seconds (5 minutes) in Streamlit. "
        "If signals were just recomputed, click Data Settings -> Sync Sectors to force "
        "st.cache_data.clear() and reload fresh data.")

    add_bullet(doc,
        "Signal version: signals.py uses SIGNAL_VERSION = '1.3'. Any change to this string "
        "triggers a full recompute of all 6,697 symbols on the next run. "
        "v1.3 entry condition: bx_monthly_dark_red==0 (green or recovering). "
        "Also fixed: indicators.py is_month_end bug that caused partial current-month BX "
        "to leak into bx_monthly_completed for the last row of the dataset.")

    add_bullet(doc,
        "Pipeline run menu: run.py in the project root provides a numbered menu. "
        "Option 8 (Force weekly indicators + signals) is the correct choice after any "
        "indicator or signal logic change. Option 9 (Vacuum) should be run monthly to "
        "reclaim disk space from old Delta table parquet files.")

    add_bullet(doc,
        "Tab 4 vs Tabs 1-3: Deep Dive recalculates all metrics live from raw signal columns. "
        "Tabs 1-3 use pre-computed bt_* snapshot values. "
        "They should match — if they differ, the snapshot may be stale.")

    add_bullet(doc,
        "Win definition (v1.2): A trade is a WIN if pct_gain > 0 (trade closed in profit). "
        "This applies regardless of whether exit was TP or SL. "
        "An SL exit at +15% counts as a win.")

    add_bullet(doc,
        "Open trades are excluded from all bt_* backtest statistics. "
        "Only completed (exited) trades contribute to win rate, profit factor, geo return, etc.")

    add_bullet(doc,
        "Sector data: 93.8% of 6,698 tickers have sector data "
        "(sourced from NASDAQ screener bulk download). "
        "The remaining 6.2% are SPACs, warrants, and rights offerings "
        "that have no sector classification.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
