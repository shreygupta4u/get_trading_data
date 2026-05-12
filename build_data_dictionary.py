"""
build_data_dictionary.py
Generates Data_Dictionary.docx for the FVB33 trading data pipeline.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_PATH = r"C:\Users\mohin\Documents\PythonCodes\get_trading_data\Data_Dictionary.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # header rows / headings
MID_BLUE   = RGBColor(0xD6, 0xE4, 0xF0)   # alt row shade
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
BLACK      = RGBColor(0x00, 0x00, 0x00)
MONO_GRAY  = RGBColor(0x1A, 0x1A, 0x2E)   # column name text

FONT_BODY  = "Calibri"
FONT_MONO  = "Courier New"


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    # remove existing shd if any
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcPr.append(borders)


def _set_col_width(table, col_widths_cm):
    """Set individual column widths (in cm) via tblGrid."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))

    # remove existing tblGrid
    for g in tbl.findall(qn("w:tblGrid")):
        tbl.remove(g)

    tblGrid = OxmlElement("w:tblGrid")
    for w_cm in col_widths_cm:
        gridCol = OxmlElement("w:gridCol")
        dxa = int(w_cm / 2.54 * 1440)
        gridCol.set(qn("w:w"), str(dxa))
        tblGrid.append(gridCol)

    # insert tblGrid after tblPr (or at start)
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _set_cell_width(cell, w_cm):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    dxa  = int(w_cm / 2.54 * 1440)
    tcW.set(qn("w:w"),    str(dxa))
    tcW.set(qn("w:type"), "dxa")
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcPr.append(tcW)


def _cell_text(cell, text, bold=False, mono=False, font_size=9,
               color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = FONT_MONO if mono else FONT_BODY
    if color:
        run.font.color.rgb = color


# ── Generic table builder ──────────────────────────────────────────────────────

def add_section_table(doc, headers, rows, col_widths_cm, note=None):
    """
    headers      : list of str
    rows         : list of list of str (same length as headers)
    col_widths_cm: list of float, must sum to ~page content width
    note         : optional italic note shown above the table
    """
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    _set_col_width(table, col_widths_cm)

    # Header row
    hdr_row = table.rows[0]
    for ci, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        _set_cell_bg(cell, "1F497D")
        _set_cell_borders(cell, "1F497D")
        _set_cell_width(cell, col_widths_cm[ci])
        _cell_text(cell, hdr, bold=True, font_size=9,
                   color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "D6E4F0" if ri % 2 == 0 else "FFFFFF"
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            _set_cell_bg(cell, fill)
            _set_cell_borders(cell, "C0C0C0", 4)
            _set_cell_width(cell, col_widths_cm[ci])
            # First column = column name → mono bold
            is_first = (ci == 0)
            _cell_text(cell, val, bold=is_first, mono=is_first,
                       font_size=9, color=MONO_GRAY if is_first else BLACK)

    doc.add_paragraph()   # spacer


# ── Heading helpers ────────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_BODY
    run.font.color.rgb = DARK_BLUE
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_BODY
    run.font.color.rgb = DARK_BLUE


def add_body(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_BODY
    run.italic = italic
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_col_count(doc, n):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {n} column{'s' if n != 1 else ''} in this section")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── Page margins ──────────────────────────────────────────────────────────────

def set_margins(doc, top=1.5, bottom=1.5, left=2.0, right=2.0):
    """Margins in cm."""
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT DATA
# ══════════════════════════════════════════════════════════════════════════════

SECTIONS = [

  ("1 — Raw OHLCV",
   "Source: yfinance (daily). Weekly OHLCV is resampled from daily bars by the orchestrator.",
   ["Column", "Daily", "Weekly", "Description", "Formula / Source"],
   [4.0, 1.0, 1.0, 5.2, 5.8],
   [
    ["symbol",      "Yes", "Yes", "Ticker symbol (e.g. AAPL)", "NASDAQ screener / yfinance"],
    ["date",        "Yes", "—",  "Trading date (daily only)",  "yfinance"],
    ["week_start",  "—",  "Yes", "Monday date of the week",    "Derived: resample('W-MON') from daily"],
    ["open",        "Yes", "Yes", "Opening price of the period", "yfinance. Weekly = first daily Open of the week"],
    ["high",        "Yes", "Yes", "Highest price of the period", "yfinance. Weekly = max(daily High) across the week"],
    ["low",         "Yes", "Yes", "Lowest price of the period",  "yfinance. Weekly = min(daily Low) across the week"],
    ["close",       "Yes", "Yes", "Closing price of the period", "yfinance. Weekly = last daily Close of the week"],
    ["volume",      "Yes", "Yes", "Total shares traded",         "yfinance. Weekly = sum(daily Volume)"],
    ["trading_days","—",  "Yes", "Number of trading days in the week", "Derived: count of daily rows resampled into that week"],
   ]),

  ("2 — Base Moving Averages",
   "All computed on Close price. Available in both Daily and Weekly tables.",
   ["Column", "Daily", "Weekly", "Description", "Formula"],
   [3.5, 1.0, 1.0, 4.5, 7.0],
   [
    ["sma_20", "Yes", "Yes", "20-bar Simple Moving Average",      "SMA(Close, 20)"],
    ["sma_33", "Yes", "Yes", "33-bar Simple Moving Average",      "SMA(Close, 33)"],
    ["ema_20", "Yes", "Yes", "20-bar Exponential Moving Average", "EMA(Close, 20), min_periods=20"],
    ["ema_33", "Yes", "Yes", "33-bar Exponential Moving Average", "EMA(Close, 33), min_periods=33"],
   ]),

  ("3 — FVB33 Fair Value Bands",
   "Source: Pine Script 'THT Fair Value Bands'. Config: FVB_PERIOD=33, FVB_SRC=OHLC4, FVB_MA_TYPE=SMA, "
   "FVB_THRESH_MODE=5 (Long SMA of TR), FVB_THRESH_WIDTH=0.4, FVB_DEV_WIDTH=3.0, FVB_COLOR_MODE=D, FVB_SMOOTH_LEN=5.",
   ["Column", "Daily", "Weekly", "Description", "Formula / Depends On"],
   [4.5, 1.0, 1.0, 4.5, 6.0],
   [
    ["fvb_src",          "Yes","Yes","Price source for basis calculation",    "(Open+High+Low+Close)/4  (OHLC4)"],
    ["fvb_basis",        "Yes","Yes","The 'fair value' centre line",           "SMA(fvb_src, 33)"],
    ["fvb_dev",          "Yes","Yes","Deviation — sets band width",            "Mode 5: SMA(TrueRange, 66).  TrueRange = max(H-L, |H-prevC|, |L-prevC|)"],
    ["fvb_upper_thresh", "Yes","Yes","Upper boundary of the inner band",       "fvb_basis + fvb_dev x 0.4"],
    ["fvb_lower_thresh", "Yes","Yes","Lower boundary of the inner band",       "fvb_basis - fvb_dev x 0.4"],
    ["fvb_upper_x1",     "Yes","Yes","X1 upper extension band",               "fvb_basis + fvb_dev x 3.4"],
    ["fvb_lower_x1",     "Yes","Yes","X1 lower extension band",               "fvb_basis - fvb_dev x 3.4"],
    ["fvb_upper_x2",     "Yes","Yes","X2 upper extension band (outermost)",   "fvb_basis + fvb_dev x 6.4"],
    ["fvb_lower_x2",     "Yes","Yes","X2 lower extension band (outermost)",   "fvb_basis - fvb_dev x 6.4"],
    ["fvb_vwap",         "Yes","Yes","VWAP proxy (per-bar)",                  "(Open+High+Low+Close)/4  (no cumulative VWAP)"],
    ["fvb_smooth_vwap",  "Yes","Yes","Smoothed VWAP",                         "EMA(fvb_vwap, 5)"],
    ["fvb_band_green",   "Yes","Yes","Bullish/Bearish flag  (1=Bull, 0=Bear)","Mode D: 1 if fvb_smooth_vwap >= fvb_basis, else 0"],
   ]),

  ("4 — FVB20 Fair Value Bands",
   "Fixed-period variant of FVB33. Hardcoded: period=20, dev=SMA(TrueRange,40), colour=EMA(Close,5) >= basis.",
   ["Column", "Daily", "Weekly", "Description", "Formula / Depends On"],
   [4.5, 1.0, 1.0, 4.5, 6.0],
   [
    ["fvb20_basis",        "Yes","Yes","FVB20 centre line",                "SMA(OHLC4, 20)"],
    ["fvb20_dev",          "Yes","Yes","FVB20 deviation",                  "SMA(TrueRange, 40)"],
    ["fvb20_upper_thresh", "Yes","Yes","FVB20 upper inner band",           "fvb20_basis + fvb20_dev x 0.4"],
    ["fvb20_lower_thresh", "Yes","Yes","FVB20 lower inner band",           "fvb20_basis - fvb20_dev x 0.4"],
    ["fvb20_upper_x1",     "Yes","Yes","FVB20 X1 upper extension",        "fvb20_basis + fvb20_dev x 3.4"],
    ["fvb20_lower_x1",     "Yes","Yes","FVB20 X1 lower extension",        "fvb20_basis - fvb20_dev x 3.4"],
    ["fvb20_upper_x2",     "Yes","Yes","FVB20 X2 upper extension",        "fvb20_basis + fvb20_dev x 6.4"],
    ["fvb20_lower_x2",     "Yes","Yes","FVB20 X2 lower extension",        "fvb20_basis - fvb20_dev x 6.4"],
    ["fvb20_band_green",   "Yes","Yes","FVB20 bullish flag  (1=Bull, 0=Bear)", "1 if EMA(Close, 5) >= fvb20_basis, else 0"],
   ]),

  ("5 — BX Trender",
   "Core formula: BX = RSI(EMA(Close,5) - EMA(Close,20), 5) - 50. "
   "Positive = bullish momentum; negative = bearish. "
   "Monthly BX is derived by resampling chart-period Close to monthly frequency. "
   "IMPORTANT: bx_monthly_completed is frozen at the previous month's value for all bars within "
   "the current (open) month — it only updates when a new month starts. "
   "This prevents the current month's partial BX from leaking into the entry signal.",
   ["Column", "Daily", "Weekly", "Description", "Formula / Depends On"],
   [4.8, 1.0, 1.0, 5.0, 5.2],
   [
    ["bx",                    "Yes","Yes","BX on the chart's own timeframe",               "RSI(EMA(Close,5) - EMA(Close,20), 5) - 50"],
    ["bx_monthly_completed",  "Yes","Yes","BX of the last fully CLOSED calendar month",
     "Monthly Close resampled -> BX -> M-1 value. Frozen throughout the current month. "
     "Updates once/month when the next month's first weekly bar arrives. "
     "Used for _bx_monthly_label display only — NOT the entry gate in signals.py v1.3."],
    ["bx_monthly_prev",       "Yes","Yes","BX of the month before last completed (M-2)",  "Same monthly series, shifted one extra month back. Used to compute bx_monthly_dark_red."],
    ["bx_monthly_dark_red",   "Yes","Yes","Entry gate AND SL trigger for FVB33 signals (v1.3)",
     "1 if bx_monthly_completed < 0 AND bx_monthly_completed <= bx_monthly_prev "
     "(last closed month is negative AND falling vs prior month). "
     "0 = either positive (green) OR negative but recovering — ENTRY ALLOWED. "
     "1 = dark red (negative and still falling) — NO ENTRY, SL if in position."],
    ["bx_monthly_realtime",   "Yes","Yes","BX of the current DEVELOPING month",           "Monthly Close resampled -> BX -> current month value (updates every bar intra-month)."],
    ["bx_realtime_dark_red",  "Yes","Yes","Dark-red flag for real-time (intra-month) BX", "1 if bx_monthly_realtime < 0 AND <= bx_monthly_completed[prev bar]. Display only."],
   ]),

  ("6 — Volume Indicators",
   "Available in both Daily and Weekly tables.",
   ["Column", "Daily", "Weekly", "Description", "Formula"],
   [3.5, 1.0, 1.0, 5.0, 6.5],
   [
    ["avg_vol_20w", "Yes","Yes","20-bar rolling average volume",      "SMA(Volume, 20)"],
    ["rel_vol",     "Yes","Yes","Relative volume vs 20-bar average",  "Volume / avg_vol_20w.  1.0 = average, 2.0 = double."],
    ["vol_up",      "Yes","Yes","Up-bar flag  (1 = up-bar)",          "1 if Close >= Open, else 0"],
   ]),

  ("7 — DBB Upper Band",
   "Available in both Daily and Weekly tables.",
   ["Column", "Daily", "Weekly", "Description", "Formula"],
   [3.5, 1.0, 1.0, 5.0, 6.5],
   [
    ["dbb_upper","Yes","Yes","Upper Double Bollinger Band",
     "SMA(Close,20) + StdDev(Close,20) x 3.0.  Used to estimate reward target at scan time."],
   ]),

  ("8 — Legacy Strategy Signals",
   "Original simple signals from indicators.py. Still stored but superseded by the sig_* columns (Section 12).",
   ["Column", "Daily", "Weekly", "Description", "Formula / Depends On"],
   [3.5, 1.0, 1.0, 5.0, 6.5],
   [
    ["entry_signal","Yes","Yes","Simple entry flag",
     "fvb_band_green==1 AND bx_monthly_dark_red==0 AND Close <= fvb_upper_thresh"],
    ["exit_signal", "Yes","Yes","Simple exit flag",
     "bx_monthly_dark_red==1"],
   ]),

  ("9 — Bull/Bear Cycle Tracker",
   "Weekly only. Stateful: enters bull when bx_monthly_dark_red==0 AND fvb_band_green==1; "
   "exits bull when bx_monthly_dark_red==1.",
   ["Column", "Daily", "Weekly", "Description", "Depends On"],
   [3.5, 1.0, 1.0, 5.5, 6.0],
   [
    ["cycle_bull",  "—","Yes","1 = currently in a bull cycle, 0 = bear/waiting",
     "bx_monthly_dark_red, fvb_band_green"],
    ["cycle_months","—","Yes","Calendar months elapsed since this cycle started",
     "cycle_bull, week_start"],
    ["cycle_pct",   "—","Yes","% price change from cycle-entry close to current close",
     "cycle_bull, Close.  Entry price = previous bar's Close when bull starts."],
   ]),

  ("10 — Bull Cycle Statistics",
   "Weekly only. Computed from ALL completed bull cycles in full history. "
   "Same scalar values are broadcast to every row for the symbol.",
   ["Column", "Daily", "Weekly", "Description", "Depends On"],
   [4.5, 1.0, 1.0, 5.5, 5.0],
   [
    ["cycle_count",        "—","Yes","Number of completed bull cycles in full history",        "cycle_bull"],
    ["cycle_win_rate",     "—","Yes","% of completed cycles ending with a gain",               "cycle_bull, cycle_pct"],
    ["cycle_geo_avg_gain", "—","Yes","Geometric mean gain per cycle — (prod(1+r))^(1/n) - 1", "All completed cycle gains"],
    ["cycle_avg_win",      "—","Yes","Arithmetic mean of winning cycle gains",                 "Completed cycles where gain > 0"],
    ["cycle_avg_loss",     "—","Yes","Arithmetic mean of losing cycle gains (negative)",        "Completed cycles where gain <= 0"],
    ["cycle_rvr",          "—","Yes","Risk-reward ratio — avg_win / abs(avg_loss)",             "cycle_avg_win, cycle_avg_loss"],
    ["cycle_profit_factor","—","Yes","sum(wins) / abs(sum(losses)) across completed cycles",   "All completed cycle gains"],
    ["cycle_avg_pullback", "—","Yes","Avg max drawdown within a bull cycle (%)",               "Low, cycle_bull.  Measures typical dip before resuming uptrend."],
    ["cycle_signal_return","—","Yes","Total compounded return following all bull cycles (%)",  "All completed cycle gains"],
   ]),

  ("11 — TPX Target Price Extensions",
   "Weekly only. Reverse-engineered from private TradingView indicator "
   "'THT Fair Value Bands - Simplified' by pdicarlo. "
   "Multiplier K=4.3 approximates the original (+/-~3% residual).",
   ["Column", "Daily", "Weekly", "Description", "Formula"],
   [3.0, 1.0, 1.0, 4.5, 7.5],
   [
    ["tpx1","—","Yes","First target price extension  (TP target)",
     "SMA(OHLC4,20) + 4.3 x EMA(TrueRange, 66)"],
    ["tpx2","—","Yes","Second target price extension",
     "SMA(OHLC4,20) + 8.6 x EMA(TrueRange, 66).  Exactly 2x the tpx1 extension above basis."],
   ]),

  ("12 — FVB33 Strategy Signals",
   "Weekly only. Generated by signals.py (v1.3). "
   "State machine: FLAT_SL(0) -> LONG(1) -> FLAT_TP(2). "
   "Entry requires bx_monthly_dark_red == 0 (last closed month is NOT dark red — "
   "either positive/green OR negative but recovering). "
   "Entry is blocked only when the last closed month is dark red (negative AND still falling).",
   ["Column", "Daily", "Weekly", "Description", "Depends On"],
   [4.5, 1.0, 1.0, 5.0, 5.5],
   [
    ["in_entry_zone",  "—","Yes","1 if price is in the FVB33 entry zone this bar",
     "Close, fvb_upper_thresh, fvb_lower_thresh.  Zone: fvb_lower x 0.75 <= Close <= fvb_upper x 1.03"],
    ["sig_entry",      "—","Yes","Entry signal — new position from FLAT_SL state",
     "fvb_band_green==1 AND bx_monthly_dark_red==0 AND in_entry_zone==1 AND tpx1 not null. "
     "dark_red==0 means last closed month BX is green or recovering (not dark red)."],
    ["sig_reentry",    "—","Yes","Re-entry after a TP exit (FLAT_TP state)",
     "Same as sig_entry conditions, but only fires from FLAT_TP state (after prior TP hit)"],
    ["sig_add",        "—","Yes","Add-to-position — entry zone re-touched while LONG",
     "Same entry conditions while sig_in_position==1.  No state change."],
    ["sig_tp",         "—","Yes","Take Profit hit — weekly High reached TPX1",
     "sig_in_position==1 AND High >= tpx1.  TP takes priority over SL on same bar."],
    ["sig_sl",         "—","Yes","Stop Loss — monthly BX turned dark red",
     "sig_in_position==1 AND bx_monthly_dark_red==1"],
    ["sig_in_position","—","Yes","1 = currently holding this trade",
     "Carry-forward from sig_entry/sig_reentry until sig_tp or sig_sl fires"],
    ["sig_entry_price","—","Yes","Entry price carried forward while in position",
     "Close on the entry/re-entry bar; NaN when not in position"],
   ]),

  ("13 — Backtest Statistics",
   "Weekly only. Computed from all completed trades per symbol; "
   "same scalar broadcast to every row. "
   "WIN definition: pct_gain > 0 (profitable close — includes SL exits still in profit).",
   ["Column", "Daily", "Weekly", "Description", "Depends On"],
   [4.5, 1.0, 1.0, 5.5, 5.0],
   [
    ["bt_trades",          "—","Yes","Total completed trades (TP + SL exits) in full history", "sig_tp, sig_sl"],
    ["bt_tp_count",        "—","Yes","Number of trades that exited via TP",                    "sig_tp"],
    ["bt_sl_count",        "—","Yes","Number of trades that exited via SL",                    "sig_sl"],
    ["bt_reentry_count",   "—","Yes","Number of trades that started as a re-entry",            "sig_reentry"],
    ["bt_win_rate",        "—","Yes","% of trades that closed in profit  (pct_gain > 0)",
     "Trade PnL.  Includes SL exits that were still profitable."],
    ["bt_avg_tp_pct",      "—","Yes","Average % gain of TP-exit trades",
     "TP exit price = tpx1 on TP bar vs entry price"],
    ["bt_avg_sl_pct",      "—","Yes","Average % gain/loss of SL-exit trades (usually negative)",
     "SL exit price = Close on SL bar vs entry price"],
    ["bt_profit_factor",   "—","Yes","sum(winning pcts) / abs(sum(losing pcts))",
     "All completed trade PnLs where pct>0 vs pct<=0"],
    ["bt_geo_return",      "—","Yes","Total compounded return — (prod(1+r) - 1) x 100",
     "All completed trade PnLs chained together"],
    ["bt_reentry_win_rate","—","Yes","Win rate for re-entry trades specifically",
     "Subset of trades where sig_reentry==1 at entry"],
    ["bt_weeks_in_trade",  "—","Yes","Total weeks sig_in_position==1 across ALL history",
     "sum(sig_in_position)"],
   ]),

  ("14 — Snapshot-Only Derived Columns",
   "These columns exist ONLY in the weekly snapshot (not stored in the Delta table). "
   "Computed at query time in db_backtest.py.",
   ["Column", "Daily", "Weekly / Snapshot", "Description", "Formula"],
   [4.5, 1.0, 1.5, 4.5, 5.5],
   [
    ["_dist_upper_pct",  "—","Snapshot","% distance of current price from FVB33 upper band",
     "(Close - fvb_upper_thresh) / fvb_upper_thresh x 100.  Negative = below band."],
    ["_upside_to_tp_pct","—","Snapshot","Upside remaining to TPX1 from current price",
     "(tpx1 - Close) / Close x 100"],
    ["_unrealised_pct",  "—","Snapshot","Unrealised P&L for currently held positions",
     "(Close - sig_entry_price) / sig_entry_price x 100.  NaN if not in position."],
    ["_fvb33_label",     "—","Snapshot","Human label for FVB33 state",
     "'Bullish' if fvb_band_green==1, else 'Bearish'"],
    ["_bx_monthly_label","—","Snapshot","Human label for completed monthly BX",
     "'Bullish' if bx_monthly_completed > 0; 'Bearish' if <= 0; '--' if null"],
    ["_bx_weekly_label", "—","Snapshot","Human label for chart-period (weekly) BX",
     "'Bullish' if bx > 0, else 'Bearish'"],
    ["_bx_mo_wk",        "—","Snapshot","Combined monthly + weekly BX label",
     "'Bull/Bull', 'Bull/Bear', 'Bear/Bull', or 'Bear/Bear'"],
    ["sector",           "—","Snapshot","Sector classification of the ticker",
     "Joined from universe DB (NASDAQ screener).  '--' for SPACs / warrants."],
   ]),
]

QUICK_REF = [
    ["OHLCV + date",                   "Yes", "Yes (resampled)"],
    ["Base MAs  (sma/ema 20/33)",       "Yes", "Yes"],
    ["FVB33 bands",                    "Yes", "Yes"],
    ["FVB20 bands",                    "Yes", "Yes"],
    ["BX Trender",                     "Yes", "Yes"],
    ["Volume indicators",              "Yes", "Yes"],
    ["DBB upper band",                 "Yes", "Yes"],
    ["Legacy entry/exit signals",      "Yes", "Yes"],
    ["Bull/bear cycle tracker",        "No",  "Yes"],
    ["Cycle statistics",               "No",  "Yes"],
    ["TPX1 / TPX2 targets",            "No",  "Yes"],
    ["FVB33 strategy signals (sig_*)", "No",  "Yes"],
    ["Backtest stats (bt_*)",          "No",  "Yes"],
    ["Snapshot derived cols (_*)",     "No",  "Snapshot only"],
]


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc, top=1.8, bottom=1.8, left=2.2, right=2.2)

    # ── Cover / title ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("FVB33 Trading Data Pipeline")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = FONT_BODY
    run.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Column Data Dictionary")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.name = FONT_BODY
    r2.font.color.rgb = DARK_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("Generated: April 2026")
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Intro ─────────────────────────────────────────────────────────────────
    add_h1(doc, "Introduction")
    add_body(doc,
        "This document describes every column stored in the three data stores of the "
        "FVB33 trading data pipeline: the Daily candles table, the Weekly candles table, "
        "and the Weekly Snapshot (one row per ticker, latest week). "
        "Raw OHLCV data is fetched from yfinance; all indicators and signals are computed "
        "in-house using Python/pandas and stored in Delta Lake format.")

    # ── Tables overview ───────────────────────────────────────────────────────
    add_h1(doc, "Tables Overview")
    add_section_table(
        doc,
        headers=["Table", "Row Granularity", "Populated By"],
        rows=[
            ["Daily",
             "1 row per ticker per trading day",
             "fetch_nasdaq_candles.py  +  technical_analysis.py --period daily"],
            ["Weekly",
             "1 row per ticker per week (Monday)",
             "Resampled from daily  +  technical_analysis.py --period weekly  +  signals.py"],
            ["Weekly Snapshot",
             "1 row per ticker (latest week only)",
             "Auto-rebuilt by signals.py after each run"],
        ],
        col_widths_cm=[3.5, 4.5, 9.0],
    )

    # ── Main sections ─────────────────────────────────────────────────────────
    for (title, note, headers, widths, rows) in SECTIONS:
        add_h2(doc, f"Section {title}")
        add_section_table(doc, headers=headers, rows=rows,
                          col_widths_cm=widths, note=note)
        add_col_count(doc, len(rows))

    # ── Quick reference ───────────────────────────────────────────────────────
    add_h1(doc, "Quick Reference — Daily vs Weekly Coverage")
    add_section_table(
        doc,
        headers=["Category", "Daily Table", "Weekly Table"],
        rows=QUICK_REF,
        col_widths_cm=[7.0, 3.5, 6.5],
    )

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
